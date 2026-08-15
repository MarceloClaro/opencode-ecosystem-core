# -*- coding: utf-8 -*-
"""Testes R422 — Blind Peer Review emulado + correções de submissão.

Requisitos (SPEC-935-R422):
1. Relatório de peer review existe com 3 pareceres e triagem consolidada.
2. Achados de correção obrigatória (IDs 6, 7, 8, 9, 11) refletidos no
   MD canônico e no TeX:
   - ID 9: RBEP = Revista Brasileira de Estudos Pedagógicos (não "de População");
   - ID 6: introdução diz "validações cruzadas por país e a análises de
     subperíodos" (não "por país e por tempo");
   - ID 7: seção 4.8 contém ressalva demográfica da matrícula bruta;
   - ID 8: Tabela 10 usa "Oriente Médio, Norte da África, Afeganistão e
     Paquistão";
   - ID 11: seção "Declarações" (conflito, financiamento, dados/código).
3. Correções informativas (IDs 1, 5, 10) aplicadas.
4. PDF recompilado sem Overfull/Underfull; DOCX regenerado e consistente.
"""

from pathlib import Path

AUDIT = Path(__file__).resolve().parent.parent / "academic/papers/arm_education_audit"
MD = AUDIT / "ARTIGO_RBEP_SUBMISSAO.md"
TEX = AUDIT / "latex" / "ARTIGO_RBEP_SUBMISSAO.tex"
REVIEW = AUDIT / "outputs" / "review" / "peer_review_r422.md"
PDF = AUDIT / "latex" / "ARTIGO_RBEP_SUBMISSAO.pdf"
LOG = AUDIT / "latex" / "ARTIGO_RBEP_SUBMISSAO.log"
DOCX = AUDIT / "outputs" / "docx" / "ARTIGO_RBEP_SUBMISSAO.docx"


class TestRelatorioReview:
    def test_relatorio_existe(self):
        assert REVIEW.exists()

    def test_tres_pareceres(self):
        texto = REVIEW.read_text(encoding="utf-8")
        for revisor in ["R1 — Econometria", "R2 — Economia", "R3 — Metodologia"]:
            assert revisor in texto

    def test_triagem_consolidada(self):
        texto = REVIEW.read_text(encoding="utf-8")
        assert "| ID | Severidade" in texto
        assert "| 12 |" in texto

    def test_sem_aprovacao_absoluta(self):
        texto = REVIEW.read_text(encoding="utf-8")
        assert "aprovado" not in texto.lower() or "não constitui aprovação" in texto.lower()
        assert "pronto para publicação" not in texto.lower()


class TestCorrecaoObrigatoria:
    def test_id9_rbep_pedagogicos_md(self):
        texto = MD.read_text(encoding="utf-8")
        assert "RBEP | Revista Brasileira de Estudos Pedagógicos" in texto
        assert "Estudos de População" not in texto

    def test_id9_rbep_pedagogicos_tex(self):
        texto = TEX.read_text(encoding="utf-8")
        assert "RBEP & Revista Brasileira de Estudos Pedagógicos" in texto
        assert "Estudos de População" not in texto

    def test_id6_validacao_subperiodos(self):
        texto = MD.read_text(encoding="utf-8")
        assert "validações cruzadas por país e a análises de subperíodos" in texto
        assert "validações cruzadas por país e por tempo" not in texto

    def test_id7_ressalva_demografica(self):
        texto = MD.read_text(encoding="utf-8")
        assert "não é corrigida pela composição etária" in texto

    def test_id8_rotulo_regional(self):
        texto_md = MD.read_text(encoding="utf-8")
        texto_tex = TEX.read_text(encoding="utf-8")
        assert "Oriente Médio, Norte da África, Afeganistão e Paquistão" in texto_md
        assert "Oriente Médio, Norte da África, Afeganistão e Paquistão" in texto_tex

    def test_id11_declaracoes_md(self):
        texto = MD.read_text(encoding="utf-8")
        assert "## Declarações" in texto
        assert "Conflito de interesses" in texto
        assert "Financiamento" in texto
        assert "Disponibilidade de dados e código" in texto

    def test_id11_declaracoes_tex(self):
        texto = TEX.read_text(encoding="utf-8")
        assert "\\section*{Declarações}" in texto
        assert "Disponibilidade de dados e código" in texto


class TestCorrecoesInformativas:
    def test_id1_arredondamento(self):
        texto = MD.read_text(encoding="utf-8")
        assert "0,7508 − 0,1465 = 0,6043" in texto

    def test_id5_proveniencia_canais(self):
        texto = MD.read_text(encoding="utf-8")
        assert "provenance_r413.json" in texto

    def test_id10_datas_download_verificacao(self):
        texto = MD.read_text(encoding="utf-8")
        assert "verificação contra a API em 14 de agosto de 2026" in texto


class TestArtefatosRegenerados:
    def test_pdf_sem_overfull(self):
        log = LOG.read_text(encoding="utf-8", errors="ignore")
        assert "Overfull" not in log
        assert "Underfull" not in log

    def test_pdf_existe(self):
        assert PDF.exists() and PDF.stat().st_size > 100_000

    def test_docx_consistente(self):
        import json
        import hashlib
        from docx import Document

        assert DOCX.exists()
        m = json.loads((AUDIT / "outputs" / "docx" / "MANIFEST.json").read_text(encoding="utf-8"))
        assert m["docx"]["sha256"] == hashlib.sha256(DOCX.read_bytes()).hexdigest()
        doc = Document(str(DOCX))
        texto = "\n".join(p.text for p in doc.paragraphs)
        # conteúdo de tabelas também (glossário A.3 fica em tabela)
        texto_tabs = " ".join(
            c.text for t in doc.tables for row in t.rows for c in row.cells
        )
        assert "Estudos Pedagógicos" in texto_tabs
        assert "Declarações" in texto or "Declarações" in texto_tabs

#!/usr/bin/env python3
"""Exportação DOCX do artigo RBEP com estilos ABNT (SPEC-935-R421).

Fluxo:
1. Gera reference.docx a partir do default do pandoc, com margens ABNT
   (esq/sup 3 cm; dir/inf 2 cm), fonte Times New Roman 12 e espaçamento 1,5.
2. Converte ARTIGO_RBEP_SUBMISSAO.md → .docx com pandoc --reference-doc.
3. Pós-processa com python-docx: garante margens/fonte/espaçamento, aplica
   bordas 'Table Grid' nas tabelas e registra sha256 no MANIFEST.

Uso: python3 scripts/export_docx_rbep.py
"""

import hashlib
import json
import shutil
import subprocess
import sys
from pathlib import Path

AUDIT = Path(__file__).resolve().parent.parent
MD = AUDIT / "ARTIGO_RBEP_SUBMISSAO.md"
OUT_DIR = AUDIT / "outputs" / "docx"
DOCX = OUT_DIR / "ARTIGO_RBEP_SUBMISSAO.docx"
REF_DOCX = OUT_DIR / "reference_abnt.docx"
MANIFEST = OUT_DIR / "MANIFEST.json"

PANDOC = shutil.which("pandoc")


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    h.update(path.read_bytes())
    return h.hexdigest()


def gerar_reference_abnt() -> None:
    """Gera reference.docx com margens/fonte ABNT a partir do default."""
    default = Path("/tmp/opencode/reference_default.docx")
    if not default.exists():
        subprocess.run(
            [PANDOC, "-o", str(default), "--print-default-data-file", "reference.docx"],
            check=True, capture_output=True, timeout=60,
        )
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(default))
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(3.0)
        sec.bottom_margin = Cm(2.0)

    def set_font(style_name: str, size_pt: float = 12, bold: bool | None = None):
        try:
            st = doc.styles[style_name]
        except KeyError:
            return
        st.font.name = "Times New Roman"
        st.font.size = Pt(size_pt)
        if bold is not None:
            st.font.bold = bold
        # garante fonte para caracteres latinos
        try:
            rpr = st.element.get_or_add_rPr()
            rfonts = rpr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
            )
            if rfonts is None:
                from docx.oxml.ns import qn
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.append(rfonts)
            from docx.oxml.ns import qn as _qn
            rfonts.set(_qn("w:ascii"), "Times New Roman")
            rfonts.set(_qn("w:hAnsi"), "Times New Roman")
            rfonts.set(_qn("w:eastAsia"), "Times New Roman")
        except Exception:
            pass

    set_font("Normal", 12)
    set_font("Body Text", 12)
    set_font("First Paragraph", 12)
    set_font("Compact", 11)
    set_font("Heading 1", 14, bold=True)
    set_font("Heading 2", 12, bold=True)
    set_font("Heading 3", 12, bold=True)
    set_font("Title", 16, bold=True)
    set_font("Table Caption", 10)
    set_font("Source Code", 10)

    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)
    doc.save(str(REF_DOCX))
    print(f"reference ABNT gerado: {REF_DOCX.name}")


def converter() -> None:
    if not PANDOC:
        sys.exit("pandoc indisponível")
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [PANDOC, str(MD), "-f", "markdown", "-t", "docx",
         "--reference-doc", str(REF_DOCX),
         "--toc", "--toc-depth=2",
         "-o", str(DOCX)],
        check=True, capture_output=True, timeout=120,
    )
    print(f"DOCX gerado: {DOCX.name}")


def pos_processar() -> None:
    from docx import Document
    from docx.shared import Cm, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document(str(DOCX))

    # margens ABNT garantidas
    for sec in doc.sections:
        sec.page_width = Cm(21.0)
        sec.page_height = Cm(29.7)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(3.0)
        sec.bottom_margin = Cm(2.0)

    # fonte/espaçamento dos estilos existentes
    for st in doc.styles:
        try:
            st.font.name = "Times New Roman"
            if st.font.size is None or st.font.size.pt > 16:
                st.font.size = Pt(12)
            st.paragraph_format.line_spacing = 1.5
            st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        except Exception:
            continue

    # bordas nas tabelas (Table Grid)
    n_tabelas = 0
    for tbl in doc.tables:
        try:
            tbl.style = doc.styles["Table Grid"]
        except Exception:
            pass
        n_tabelas += 1

    doc.save(str(DOCX))
    print(f"pós-processamento OK: {n_tabelas} tabelas com bordas")


def manifest() -> None:
    dados = {
        "artigo_md": {
            "sha256": sha256(MD),
            "caminho": str(MD.relative_to(AUDIT)),
        },
        "docx": {
            "sha256": sha256(DOCX),
            "caminho": str(DOCX.relative_to(AUDIT)),
            "bytes": DOCX.stat().st_size,
        },
        "reference": {
            "sha256": sha256(REF_DOCX),
            "caminho": str(REF_DOCX.relative_to(AUDIT)),
        },
        "pandoc": PANDOC,
        "data": "2026-08-14",
        "nota": "Exportação DOCX com estilos ABNT (margens 3/3/2/2 cm, Times New Roman 12, espaçamento 1,5).",
    }
    MANIFEST.write_text(json.dumps(dados, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"MANIFEST: {MANIFEST.name}")


def main() -> int:
    gerar_reference_abnt()
    converter()
    pos_processar()
    manifest()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
